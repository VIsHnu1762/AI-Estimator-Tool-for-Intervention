"""
Document text extraction service
Handles PDF, DOCX, images, and plain text
"""
import logging
from pathlib import Path
from typing import Optional
import PyPDF2
import docx
from PIL import Image
import pytesseract
from pdf2image import convert_from_path

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        # Set tesseract path
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from document based on file type
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif extension == '.docx':
                return await self._extract_from_docx(file_path)
            elif extension in ['.txt']:
                return await self._extract_from_text(file_path)
            elif extension in ['.png', '.jpg', '.jpeg']:
                return await self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}", exc_info=True)
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        
        try:
            # Try direct text extraction first
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If no text extracted, use OCR
            if len(text.strip()) < 100:
                logger.info(f"PDF appears to be scanned. Using OCR...")
                text = await self._ocr_pdf(file_path)
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            # Fallback to OCR
            return await self._ocr_pdf(file_path)
    
    async def _ocr_pdf(self, file_path: str) -> str:
        """Perform OCR on PDF"""
        text = ""
        try:
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300)
            
            # OCR each page
            for i, image in enumerate(images):
                logger.info(f"OCR processing page {i+1}/{len(images)}")
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += page_text + "\n"
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error in OCR: {str(e)}")
            raise
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            raise
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
    
    async def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='eng')
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error extracting from image: {str(e)}")
            raise
